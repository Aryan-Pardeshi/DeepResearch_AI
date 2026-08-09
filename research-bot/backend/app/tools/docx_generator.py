import logging
from pathlib import Path
from typing import Dict, Any, List
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

logger = logging.getLogger(__name__)

def _set_cell_background(cell, fill_hex: str):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_paper_docx(state: Dict[str, Any], output_path: str) -> str:
    """Generates a styled academic Word document (.docx) from ResearchModeState."""
    doc = docx.Document()

    # Set 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Title
    title_text = state.get("title") or state.get("problem_statement") or "Research Report"
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)
    title_p.paragraph_format.space_after = Pt(12)

    # Abstract Callout Box
    abstract_text = state.get("abstract", "")
    if isinstance(abstract_text, str) and abstract_text.strip():
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        _set_cell_background(cell, "F0F4F8")
        
        ap = cell.paragraphs[0]
        arun_bold = ap.add_run("ABSTRACT\n")
        arun_bold.bold = True
        arun_bold.font.size = Pt(10)
        arun_bold.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)
        
        arun_text = ap.add_run(abstract_text.strip())
        arun_text.font.size = Pt(10.5)
        arun_text.font.italic = True
        ap.paragraph_format.space_after = Pt(4)

        # Keywords
        keywords = state.get("keywords", [])
        if keywords:
            kw_str = ", ".join(keywords) if isinstance(keywords, list) else str(keywords)
            kw_p = cell.add_paragraph()
            kw_bold = kw_p.add_run("Keywords: ")
            kw_bold.bold = True
            kw_bold.font.size = Pt(9.5)
            kw_text = kw_p.add_run(kw_str)
            kw_text.font.size = Pt(9.5)
            kw_text.font.italic = True

        doc.add_paragraph() # Spacer

    # Helper for adding formatted sections
    def add_section(heading_title: str, content: Any, is_list: bool = False):
        if not content:
            return
        h = doc.add_heading(level=1)
        hrun = h.add_run(heading_title)
        hrun.font.name = 'Times New Roman'
        hrun.font.size = Pt(14)
        hrun.font.bold = True
        hrun.font.color.rgb = RGBColor(0x1A, 0x2B, 0x4C)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

        if is_list and isinstance(content, list):
            for item in content:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(str(item))
        elif isinstance(content, list):
            for item in content:
                p = doc.add_paragraph()
                p.add_run(str(item))
        else:
            p = doc.add_paragraph()
            p.add_run(str(content))

    # Core Academic Sections
    add_section("1. Introduction", state.get("introduction"))
    add_section("2. Literature Review", state.get("literature_review"))
    add_section("3. Research Gap", state.get("research_gap"))
    add_section("4. Research Objectives", state.get("research_objectives"), is_list=True)
    add_section("5. Research Questions", state.get("research_questions"), is_list=True)
    add_section("6. Conceptual Framework", state.get("conceptual_framework"))
    add_section("7. Hypotheses", state.get("hypotheses"), is_list=True)
    add_section("8. Research Design", state.get("research_design"))
    add_section("9. Data Collection Plan", state.get("data_collection_plan"))
    add_section("10. Data Analysis Plan", state.get("data_analysis_plan"))
    add_section("11. Results", state.get("results"))
    add_section("12. Discussion", state.get("discussion"))
    add_section("13. Practical & Theoretical Implications", state.get("implications"))
    add_section("14. Limitations", state.get("limitations"))
    add_section("15. Conclusion", state.get("conclusion"))
    add_section("16. Future Scope", state.get("future_scope"))
    add_section("17. References", state.get("references"), is_list=True)
    add_section("18. Appendices", state.get("appendices"), is_list=True)

    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    logger.info(f"Generated DOCX report at: {out_path}")
    return str(out_path)
